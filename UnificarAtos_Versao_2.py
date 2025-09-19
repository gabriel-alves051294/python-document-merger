# -*- coding: utf-8 -*-
"""
Versão Definitiva: Automação do Processamento de Atos Normativos.

Este script executa um fluxo completo e otimizado:
1.  Filtra arquivos .doc e .docx, descartando os "revogados"
    (com mais de 90% do texto tachado).
2.  Converte os arquivos válidos para .docx, mantendo a estrutura de pastas.
3.  Extrai o texto limpo, implementando a lógica correta para hyperlinks:
    - O TEXTO de um link é preservado.
    - O URL de um link é descartado.
    - Textos com tachado são removidos.
4.  Consolida o texto de cada categoria em arquivos .txt.
5.  Formata a saída .txt com tags XML e metadados para máxima compatibilidade
    e precisão com LLMs (ex: NotebookLM), evitando a mistura de conteúdo.
6.  Garante que nenhum ato seja dividido ao meio durante a fragmentação dos
    arquivos .txt.
"""

import os
import subprocess
import shutil
import logging
from tqdm import tqdm
import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import _Cell, Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.text.hyperlink import CT_Hyperlink
from docx.oxml.ns import qn
from collections import defaultdict
from typing import List, Optional, Dict
from datetime import datetime

# --- CONFIGURAÇÕES ---
# 1. Estrutura de Diretórios
PASTA_ENTRADA = r'C:\ProcessarAtos\Entrada'
PASTA_SAIDA_DOCX = r'C:\ProcessarAtos\Saida_DOCX'
PASTA_SAIDA_TXT = r'C:\ProcessarAtos\Saida_TXT'
ARQUIVO_DE_LOG = r'C:\ProcessarAtos\log_processamento.log'

# 2. Executáveis
CAMINHO_SOFFICE = r'C:\Program Files\LibreOffice\program\soffice.exe'

# 3. Regras de Processamento
PERCENTUAL_MINIMO_TACHADO_PARA_IGNORAR: float = 80.0
MAX_TAMANHO_TXT_MB: int = 2
MAX_TAMANHO_TXT_BYTES: int = MAX_TAMANHO_TXT_MB * 1024 * 1024


# --- FIM DAS CONFIGURAÇÕES ---


def setup_logging():
    """Configura o sistema de logging para registrar eventos em arquivo e no console."""
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(ARQUIVO_DE_LOG, mode='w', encoding='utf-8'),
            logging.StreamHandler()
        ]
    )


def converter_doc_para_docx(doc_path: str, output_dir_especifico: Optional[str] = None) -> Optional[str]:
    """Usa o LibreOffice para converter .doc para .docx."""
    try:
        output_dir = output_dir_especifico or os.path.join(os.path.dirname(doc_path), 'convertidos_temp')
        os.makedirs(output_dir, exist_ok=True)

        cmd = [CAMINHO_SOFFICE, '--headless', '--convert-to', 'docx', '--outdir', output_dir, doc_path]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120)

        base_name = os.path.basename(doc_path)
        new_docx_path = os.path.join(output_dir, os.path.splitext(base_name)[0] + '.docx')

        if os.path.exists(new_docx_path):
            return new_docx_path
        else:
            raise FileNotFoundError("Arquivo convertido não foi encontrado após a execução do LibreOffice.")

    except FileNotFoundError:
        logging.error(f"O executável 'soffice.exe' não foi encontrado em: {CAMINHO_SOFFICE}.")
    except subprocess.TimeoutExpired:
        logging.error(f"Timeout ao converter '{os.path.basename(doc_path)}'. O processo demorou mais de 120s.")
    except subprocess.CalledProcessError as e:
        logging.error(
            f"Erro do LibreOffice ao converter '{os.path.basename(doc_path)}': {e.stderr.decode('utf-8', errors='ignore')}")
    except Exception as e:
        logging.error(f"Erro inesperado ao converter '{os.path.basename(doc_path)}': {e}", exc_info=True)

    return None


def iter_block_items(parent):
    """Iterador que produz parágrafos e tabelas na ordem correta do documento."""
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Tipo de 'parent' não suportado")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def analisar_percentual_tachado(docx_path: str) -> float:
    """Analisa um arquivo .docx e retorna o percentual de caracteres tachados."""
    try:
        documento = docx.Document(docx_path)
        total_caracteres, caracteres_tachados = 0, 0

        def contar_runs(runs):
            nonlocal total_caracteres, caracteres_tachados
            for run in runs:
                num_chars = len(run.text)
                total_caracteres += num_chars
                if run.font.strike: caracteres_tachados += num_chars

        for block in iter_block_items(documento):
            if isinstance(block, Paragraph):
                contar_runs(block.runs)
            elif isinstance(block, Table):
                for row in block.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs: contar_runs(p.runs)

        return (caracteres_tachados / total_caracteres) * 100 if total_caracteres > 0 else 0.0
    except Exception as e:
        logging.error(f"Erro ao analisar '{os.path.basename(docx_path)}': {e}", exc_info=True)
        return -1.0


def extrair_texto_limpo_de_docx(docx_path: str) -> Optional[str]:
    """
    Extrai texto de um .docx, ignorando trechos tachados mas preservando
    APENAS O TEXTO dos hyperlinks, descartando o URL.
    """
    try:
        documento = docx.Document(docx_path)
        texto_completo = []

        def processar_paragrafo(p: Paragraph) -> str:
            """Processa um único parágrafo, tratando hyperlinks e runs normais."""
            texto_paragrafo = ""
            for item in p._p.iterchildren():
                # Se o item for um hyperlink
                if isinstance(item, CT_Hyperlink):
                    # Extrai o texto visível do link
                    texto_link = "".join(t.text for r in item.r_lst for t in r.t_lst)
                    # Adiciona APENAS o texto visível ao resultado
                    texto_paragrafo += texto_link

                # Se for um 'run' de texto normal
                elif isinstance(item, docx.oxml.text.run.CT_R):
                    run = Run(item, p)
                    # Adiciona o texto apenas se não estiver tachado
                    if not run.font.strike:
                        texto_paragrafo += run.text
            return texto_paragrafo

        # Itera sobre todos os blocos (parágrafos e tabelas) do documento
        for block in iter_block_items(documento):
            if isinstance(block, Paragraph):
                texto_completo.append(processar_paragrafo(block))
            elif isinstance(block, Table):
                for row in block.rows:
                    celulas = ["\n".join([processar_paragrafo(p) for p in cell.paragraphs]) for cell in row.cells]
                    texto_completo.append("\t".join(celulas))

        return "\n".join(filter(None, texto_completo))

    except Exception as e:
        logging.error(f"Erro ao extrair texto de '{os.path.basename(docx_path)}': {e}", exc_info=True)
        return None


def etapa_1_filtrar_revogados(todos_arquivos: List[str]) -> List[str]:
    """Filtra a lista de arquivos, retornando apenas os que não são revogados."""
    logging.info("--- ETAPA 1: Analisando e filtrando arquivos revogados ---")
    arquivos_validos = []
    for file_path in tqdm(todos_arquivos, desc="Analisando arquivos"):
        path_analise = file_path
        is_doc = file_path.lower().endswith('.doc')

        if is_doc: path_analise = converter_doc_para_docx(file_path)

        if path_analise:
            percentual = analisar_percentual_tachado(path_analise)
            if 0 <= percentual < PERCENTUAL_MINIMO_TACHADO_PARA_IGNORAR:
                arquivos_validos.append(file_path)
            else:
                logging.warning(f"IGNORADO (REVOGADO): '{os.path.basename(file_path)}' ({percentual:.2f}% tachado)")
            if is_doc and os.path.exists(os.path.dirname(path_analise)):
                shutil.rmtree(os.path.dirname(path_analise), ignore_errors=True)
    return arquivos_validos


def etapa_2_converter_validos(arquivos_validos: List[str]) -> Dict[str, List[str]]:
    """Converte arquivos válidos para .docx e os organiza por categoria."""
    logging.info(f"--- ETAPA 2: Convertendo {len(arquivos_validos)} arquivo(s) válido(s) para a pasta Saida_DOCX ---")
    categorias = defaultdict(list)
    for file_path in tqdm(arquivos_validos, desc="Convertendo para .docx"):
        try:
            relative_path = os.path.relpath(os.path.dirname(file_path), PASTA_ENTRADA)
            dest_folder = os.path.join(PASTA_SAIDA_DOCX, relative_path)
            os.makedirs(dest_folder, exist_ok=True)

            nome_sem_ext = os.path.splitext(os.path.basename(file_path))[0]
            caminho_final_docx = os.path.join(dest_folder, f"{nome_sem_ext}.docx")

            if file_path.lower().endswith('.docx'):
                shutil.copy(file_path, caminho_final_docx)
            else:
                converter_doc_para_docx(file_path, dest_folder)

            if os.path.exists(caminho_final_docx):
                categorias[relative_path].append(caminho_final_docx)
            else:
                logging.error(f"Falha ao gerar o arquivo .docx final para: {file_path}")
        except Exception as e:
            logging.error(f"Erro ao processar a conversão de {file_path}: {e}", exc_info=True)
    return categorias


def etapa_3_4_consolidar_e_gerar_txt(categorias: Dict[str, List[str]]):
    """Consolida e gera arquivos .txt para cada categoria com formato otimizado para LLMs."""
    logging.info("--- ETAPA 3/4: Consolidando e gerando arquivos .txt por categoria ---")
    for relative_path, lista_docx in tqdm(categorias.items(), desc="Processando categorias"):
        dest_folder_txt = os.path.join(PASTA_SAIDA_TXT, relative_path)
        os.makedirs(dest_folder_txt, exist_ok=True)

        nome_categoria = os.path.basename(relative_path).lower().replace(" ", "_") if relative_path != '.' else 'raiz'
        nome_base_txt = f"{nome_categoria}"

        bloco_texto_categoria = []
        for docx_path in lista_docx:
            texto_limpo = extrair_texto_limpo_de_docx(docx_path)
            if texto_limpo and texto_limpo.strip():
                nome_original = os.path.basename(docx_path)

                # --- Bloco de formatação otimizado para LLMs ---
                data_hora_atual = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                categoria_atual = os.path.basename(relative_path) if relative_path != '.' else 'Raiz'

                bloco_formatado = f"""<documento fonte="{nome_original}">
<metadados>
  <categoria>{categoria_atual}</categoria>
  <arquivo_original>{nome_original}</arquivo_original>
  <data_processamento>{data_hora_atual}</data_processamento>
</metadados>
<conteudo>
# Ato Normativo: {nome_original}

{texto_limpo}
</conteudo>
</documento>
"""
                bloco_texto_categoria.append(bloco_formatado + "\n\n")

        if not bloco_texto_categoria: continue

        arquivo_txt_indice = 1
        nome_arquivo_atual = os.path.join(dest_folder_txt, f"{nome_base_txt}_{arquivo_txt_indice:02d}.txt")
        f_txt = open(nome_arquivo_atual, 'w', encoding='utf-8')

        try:
            for conteudo_doc in bloco_texto_categoria:
                tamanho_bytes = len(conteudo_doc.encode('utf-8'))
                # Lógica para não dividir um documento no meio
                if f_txt.tell() + tamanho_bytes > MAX_TAMANHO_TXT_BYTES and f_txt.tell() > 0:
                    f_txt.close()
                    arquivo_txt_indice += 1
                    nome_arquivo_atual = os.path.join(dest_folder_txt, f"{nome_base_txt}_{arquivo_txt_indice:02d}.txt")
                    f_txt = open(nome_arquivo_atual, 'w', encoding='utf-8')
                f_txt.write(conteudo_doc)
        finally:
            f_txt.close()


def main():
    """Função principal que orquestra todo o processo."""
    setup_logging()
    logging.info(">>> INICIANDO PROCESSO DE CLASSIFICAÇÃO, CONVERSÃO E CONSOLIDAÇÃO DE ATOS NORMATIVOS <<<")

    os.makedirs(PASTA_SAIDA_DOCX, exist_ok=True)
    os.makedirs(PASTA_SAIDA_TXT, exist_ok=True)

    try:
        todos_arquivos = [os.path.join(r, f) for r, _, fs in os.walk(PASTA_ENTRADA) for f in fs if
                          f.lower().endswith(('.doc', '.docx')) and not f.startswith('~')]
        if not todos_arquivos:
            logging.warning(f"Nenhum arquivo .doc ou .docx encontrado em '{PASTA_ENTRADA}'.")
            return

        arquivos_validos = etapa_1_filtrar_revogados(todos_arquivos)
        categorias = etapa_2_converter_validos(arquivos_validos)
        etapa_3_4_consolidar_e_gerar_txt(categorias)

        logging.info("--- PROCESSO CONCLUÍDO COM SUCESSO ---")
        logging.info(f"Arquivos válidos convertidos para .docx salvos em: '{PASTA_SAIDA_DOCX}'")
        logging.info(f"Textos consolidados e otimizados para IA salvos em: '{PASTA_SAIDA_TXT}'")

    except Exception as e:
        logging.critical(f"Ocorreu um erro fatal no processo: {e}", exc_info=True)
    finally:
        logging.info(">>> FIM DA EXECUÇÃO <<<")


if __name__ == "__main__":
    main()
